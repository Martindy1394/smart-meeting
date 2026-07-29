"""Multi-format meeting export (TXT / DOCX / PDF).

Bundles original transcript, English translation, structured summary, and
timestamped segments when available — see docs/PRODUCT.md.
"""
from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any

from docx import Document
from fpdf import FPDF


def _safe_filename(title: str, fmt: str) -> str:
    base = (title or "meeting").strip() or "meeting"
    base = re.sub(r"[^\w\-]+", "_", base).strip("_") or "meeting"
    base = re.sub(r"_+", "_", base)[:80]
    return f"{base}_smart_meeting.{fmt}"


def _fmt_ts(seconds: float | None) -> str:
    if seconds is None:
        return ""
    try:
        s = max(0.0, float(seconds))
    except (TypeError, ValueError):
        return ""
    m, sec = divmod(int(s), 60)
    h, m = divmod(m, 60)
    if h:
        return f"{h:d}:{m:02d}:{sec:02d}"
    return f"{m:d}:{sec:02d}"


def build_export_sections(meeting: Any) -> list[tuple[str, str]]:
    """Return ordered (heading, body) sections for export."""
    sections: list[tuple[str, str]] = []

    meta_lines = [
        f"Title: {(meeting.title or 'Untitled meeting').strip()}",
        f"Venue: {(meeting.venue or '').strip() or '—'}",
        f"Language: {(meeting.language or 'auto').strip()}",
        f"Status: {(meeting.status or '').strip() or '—'}",
    ]
    if getattr(meeting, "meeting_date", None):
        md = meeting.meeting_date
        if isinstance(md, datetime):
            meta_lines.append(f"Meeting date: {md.isoformat()}")
    sections.append(("Meeting details", "\n".join(meta_lines)))

    transcript = (getattr(meeting, "final_transcript", None) or "").strip()
    sections.append(
        (
            "Verbatim transcript (original language)",
            transcript or "(No transcript yet.)",
        )
    )

    segments = list(getattr(meeting, "segments", None) or [])
    if segments:
        lines: list[str] = []
        for seg in segments:
            text = (getattr(seg, "text", None) or "").strip()
            if not text:
                continue
            start = _fmt_ts(getattr(seg, "start_time", None))
            end = _fmt_ts(getattr(seg, "end_time", None))
            stamp = f"[{start}–{end}] " if start or end else ""
            lines.append(f"{stamp}{text}")
        if lines:
            sections.append(("Timestamped segments", "\n".join(lines)))

    translation = (getattr(meeting, "translation", None) or "").strip()
    tlang = (getattr(meeting, "translation_language", None) or "English").strip()
    sections.append(
        (
            f"English translation ({tlang or 'English'})",
            translation or "(No English translation yet.)",
        )
    )

    summary = (getattr(meeting, "summary", None) or "").strip()
    sfmt = (getattr(meeting, "summary_format", None) or "").strip()
    heading = "Structured summary"
    if sfmt:
        heading = f"Structured summary ({sfmt})"
    sections.append((heading, summary or "(No summary yet.)"))

    return sections


def render_txt(meeting: Any) -> bytes:
    parts: list[str] = ["Smart Meeting export", ""]
    for heading, body in build_export_sections(meeting):
        parts.append(heading)
        parts.append("=" * len(heading))
        parts.append(body)
        parts.append("")
    return "\n".join(parts).encode("utf-8")


def render_docx(meeting: Any) -> bytes:
    doc = Document()
    doc.add_heading("Smart Meeting export", level=0)
    for heading, body in build_export_sections(meeting):
        doc.add_heading(heading, level=1)
        for para in body.split("\n"):
            doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class _MeetingPDF(FPDF):
    def footer(self) -> None:
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(120, 120, 120)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")


def _pdf_safe(text: str) -> str:
    """FPDF core fonts are Latin-1; replace unsupported chars."""
    cleaned = (
        (text or "")
        .replace("\u2022", "-")
        .replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
    )
    return cleaned.encode("latin-1", errors="replace").decode("latin-1")


def render_pdf(meeting: Any) -> bytes:
    pdf = _MeetingPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, _pdf_safe("Smart Meeting export"))
    pdf.ln(2)
    for heading, body in build_export_sections(meeting):
        pdf.set_font("Helvetica", "B", 12)
        pdf.multi_cell(0, 8, _pdf_safe(heading))
        pdf.ln(1)
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5, _pdf_safe(body))
        pdf.ln(4)
    out = pdf.output()
    if isinstance(out, (bytes, bytearray)):
        return bytes(out)
    return str(out).encode("latin-1", errors="replace")


def export_meeting(meeting: Any, fmt: str) -> tuple[bytes, str, str]:
    """Return ``(payload, media_type, filename)`` for ``fmt`` in txt|docx|pdf."""
    kind = (fmt or "txt").strip().lower()
    if kind not in {"txt", "docx", "pdf"}:
        raise ValueError("format must be txt, docx, or pdf")
    if kind == "txt":
        data = render_txt(meeting)
        media = "text/plain; charset=utf-8"
    elif kind == "docx":
        data = render_docx(meeting)
        media = (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        )
    else:
        data = render_pdf(meeting)
        media = "application/pdf"
    return data, media, _safe_filename(getattr(meeting, "title", "") or "", kind)
